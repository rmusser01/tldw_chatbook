# export_manager.py
# Description: Export manager for briefings and subscription content
#
# This module handles exporting briefings to various formats:
# - PDF (using ReportLab)
# - EPUB (for e-readers)
# - DOCX (for editing)
# - Enhanced HTML
# - Audio (via TTS)
#
# Imports
import os
import json
from pathlib import Path
from datetime import datetime, timezone
from typing import Optional, Dict, Any, List, Union
import tempfile
import mimetypes
#
# Third-Party Imports
from loguru import logger
try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.platypus import Table, TableStyle, Image
    from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False
    logger.warning("ReportLab not installed. PDF export disabled.")

try:
    from ebooklib import epub
    HAS_EBOOKLIB = True
except ImportError:
    HAS_EBOOKLIB = False
    logger.warning("ebooklib not installed. EPUB export disabled.")

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.warning("python-docx not installed. DOCX export disabled.")
#
# Local Imports
from ..Metrics.metrics_logger import log_counter, log_histogram
#
########################################################################################################################
#
# Export Manager
#
########################################################################################################################

class ExportManager:
    """Manages exporting briefings to various formats."""
    
    def __init__(self, output_dir: Optional[Path] = None):
        """
        Initialize export manager.
        
        Args:
            output_dir: Default output directory for exports
        """
        self.output_dir = output_dir or Path.home() / "Downloads" / "briefings"
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    async def export_briefing(self,
                            content: str,
                            format: str,
                            metadata: Dict[str, Any],
                            output_path: Optional[Path] = None) -> Path:
        """
        Export briefing to specified format.
        
        Args:
            content: Briefing content (markdown format)
            format: Export format (pdf, epub, docx, html)
            metadata: Briefing metadata (title, date, etc.)
            output_path: Optional specific output path
            
        Returns:
            Path to exported file
        """
        # Generate default filename if not provided
        if not output_path:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{metadata.get('name', 'briefing')}_{timestamp}.{format}"
            output_path = self.output_dir / filename
        
        # Ensure parent directory exists
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Export based on format
        if format == 'pdf':
            return await self._export_pdf(content, metadata, output_path)
        elif format == 'epub':
            return await self._export_epub(content, metadata, output_path)
        elif format == 'docx':
            return await self._export_docx(content, metadata, output_path)
        elif format == 'html':
            return await self._export_html(content, metadata, output_path)
        else:
            raise ValueError(f"Unsupported export format: {format}")
    
    async def _export_pdf(self, content: str, metadata: Dict[str, Any], output_path: Path) -> Path:
        """Export to PDF format."""
        if not HAS_REPORTLAB:
            raise ImportError("ReportLab is required for PDF export. Install with: pip install reportlab")
        
        start_time = datetime.now()
        
        try:
            # Create PDF document
            doc = SimpleDocTemplate(
                str(output_path),
                pagesize=letter,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=18
            )
            
            # Container for flowables
            story = []
            
            # Get styles
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Title'],
                fontSize=24,
                textColor=colors.HexColor('#1a1a1a'),
                spaceAfter=30,
                alignment=TA_CENTER
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading1'],
                fontSize=16,
                textColor=colors.HexColor('#2c3e50'),
                spaceAfter=12,
                spaceBefore=24
            )
            
            subheading_style = ParagraphStyle(
                'CustomSubheading',
                parent=styles['Heading2'],
                fontSize=14,
                textColor=colors.HexColor('#34495e'),
                spaceAfter=10,
                spaceBefore=18
            )
            
            body_style = ParagraphStyle(
                'CustomBody',
                parent=styles['BodyText'],
                fontSize=11,
                alignment=TA_JUSTIFY,
                spaceAfter=12
            )
            
            # Add title
            title = metadata.get('name', 'Briefing')
            story.append(Paragraph(title, title_style))
            
            # Add metadata
            date_str = metadata.get('generated_at', datetime.now().isoformat())
            story.append(Paragraph(f"Generated: {date_str}", styles['Normal']))
            story.append(Spacer(1, 0.5*inch))
            
            # Parse markdown content
            lines = content.split('\n')
            
            for line in lines:
                line = line.strip()
                
                if not line:
                    story.append(Spacer(1, 0.2*inch))
                    continue
                
                # Headers
                if line.startswith('# '):
                    text = line[2:].strip()
                    story.append(Paragraph(text, heading_style))
                elif line.startswith('## '):
                    text = line[3:].strip()
                    story.append(Paragraph(text, subheading_style))
                elif line.startswith('### '):
                    text = line[4:].strip()
                    story.append(Paragraph(text, styles['Heading3']))
                
                # Lists
                elif line.startswith('- ') or line.startswith('* '):
                    text = '• ' + line[2:].strip()
                    story.append(Paragraph(text, body_style))
                elif line[0].isdigit() and line[1:3] == '. ':
                    story.append(Paragraph(line, body_style))
                
                # Regular text
                else:
                    # Convert markdown formatting
                    text = self._convert_markdown_for_pdf(line)
                    story.append(Paragraph(text, body_style))
            
            # Add footer
            story.append(Spacer(1, 0.5*inch))
            story.append(Paragraph("---", styles['Normal']))
            
            footer_text = f"This briefing contains {metadata.get('item_count', 0)} items from {metadata.get('source_count', 0)} sources."
            story.append(Paragraph(footer_text, styles['Italic']))
            
            # Build PDF
            doc.build(story)
            
            # Log metrics
            duration = (datetime.now() - start_time).total_seconds()
            log_histogram("export_duration", duration, labels={"format": "pdf"})
            log_counter("exports_created", labels={"format": "pdf"})
            
            logger.info(f"Exported PDF to: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error exporting PDF: {str(e)}")
            raise
    
    async def _export_epub(self, content: str, metadata: Dict[str, Any], output_path: Path) -> Path:
        """Export to EPUB format."""
        if not HAS_EBOOKLIB:
            raise ImportError("ebooklib is required for EPUB export. Install with: pip install ebooklib")
        
        try:
            # Create EPUB book
            book = epub.EpubBook()
            
            # Set metadata
            book.set_identifier(f"briefing_{datetime.now().timestamp()}")
            book.set_title(metadata.get('name', 'Briefing'))
            book.set_language('en')
            
            book.add_author('Briefing Generator')
            
            # Convert markdown to HTML
            import markdown
            html_content = markdown.markdown(content, extensions=['extra'])
            
            # Create chapter
            chapter = epub.EpubHtml(
                title=metadata.get('name', 'Briefing'),
                file_name='briefing.xhtml',
                lang='en'
            )
            
            # Wrap content in proper HTML
            chapter_html = f"""
            <html>
            <head>
                <title>{metadata.get('name', 'Briefing')}</title>
            </head>
            <body>
                <h1>{metadata.get('name', 'Briefing')}</h1>
                <p><em>Generated: {metadata.get('generated_at', datetime.now().isoformat())}</em></p>
                {html_content}
                <hr/>
                <p><em>This briefing contains {metadata.get('item_count', 0)} items from {metadata.get('source_count', 0)} sources.</em></p>
            </body>
            </html>
            """
            
            chapter.content = chapter_html
            
            # Add chapter to book
            book.add_item(chapter)
            
            # Define Table of Contents
            book.toc = (epub.Link('briefing.xhtml', metadata.get('name', 'Briefing'), 'briefing'),)
            
            # Add navigation files
            book.add_item(epub.EpubNcx())
            book.add_item(epub.EpubNav())
            
            # Define spine
            book.spine = ['nav', chapter]
            
            # Write EPUB file
            epub.write_epub(str(output_path), book, {})
            
            logger.info(f"Exported EPUB to: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error exporting EPUB: {str(e)}")
            raise
    
    async def _export_docx(self, content: str, metadata: Dict[str, Any], output_path: Path) -> Path:
        """Export to DOCX format."""
        if not HAS_DOCX:
            raise ImportError("python-docx is required for DOCX export. Install with: pip install python-docx")
        
        try:
            # Create document
            doc = Document()
            
            # Add title
            title = doc.add_heading(metadata.get('name', 'Briefing'), 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add metadata
            p = doc.add_paragraph()
            p.add_run(f"Generated: {metadata.get('generated_at', datetime.now().isoformat())}").italic = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add line break
            doc.add_paragraph()
            
            # Parse markdown content
            lines = content.split('\n')
            
            for line in lines:
                line = line.strip()
                
                if not line:
                    doc.add_paragraph()
                    continue
                
                # Headers
                if line.startswith('# '):
                    doc.add_heading(line[2:].strip(), level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:].strip(), level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:].strip(), level=3)
                
                # Lists
                elif line.startswith('- ') or line.startswith('* '):
                    p = doc.add_paragraph(line[2:].strip(), style='List Bullet')
                elif line[0].isdigit() and line[1:3] == '. ':
                    p = doc.add_paragraph(line[3:].strip(), style='List Number')
                
                # Regular text
                else:
                    p = doc.add_paragraph()
                    self._add_markdown_to_paragraph(p, line)
            
            # Add footer
            doc.add_page_break()
            footer = doc.add_paragraph()
            footer.add_run("---").bold = True
            footer = doc.add_paragraph()
            footer.add_run(f"This briefing contains {metadata.get('item_count', 0)} items from {metadata.get('source_count', 0)} sources.").italic = True
            
            # Save document
            doc.save(str(output_path))
            
            logger.info(f"Exported DOCX to: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error exporting DOCX: {str(e)}")
            raise
    
    async def _export_html(self, content: str, metadata: Dict[str, Any], output_path: Path) -> Path:
        """Export to enhanced HTML format."""
        try:
            import markdown
            
            # Convert markdown to HTML
            html_body = markdown.markdown(content, extensions=['extra', 'codehilite', 'toc'])
            
            # Create full HTML document
            html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{metadata.get('name', 'Briefing')}</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            background-color: #f5f5f5;
        }}
        .container {{
            background-color: white;
            padding: 40px;
            border-radius: 10px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }}
        h1, h2, h3 {{
            color: #2c3e50;
            margin-top: 30px;
        }}
        h1 {{
            text-align: center;
            font-size: 2.5em;
            margin-bottom: 10px;
        }}
        .metadata {{
            text-align: center;
            color: #7f8c8d;
            margin-bottom: 40px;
        }}
        a {{
            color: #3498db;
            text-decoration: none;
        }}
        a:hover {{
            text-decoration: underline;
        }}
        code {{
            background-color: #f4f4f4;
            padding: 2px 4px;
            border-radius: 3px;
            font-family: 'Consolas', 'Monaco', monospace;
        }}
        pre {{
            background-color: #f4f4f4;
            padding: 15px;
            border-radius: 5px;
            overflow-x: auto;
        }}
        blockquote {{
            border-left: 4px solid #3498db;
            margin: 20px 0;
            padding-left: 20px;
            color: #555;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 20px 0;
        }}
        th, td {{
            border: 1px solid #ddd;
            padding: 12px;
            text-align: left;
        }}
        th {{
            background-color: #f4f4f4;
            font-weight: bold;
        }}
        .footer {{
            margin-top: 50px;
            padding-top: 20px;
            border-top: 1px solid #eee;
            color: #7f8c8d;
            text-align: center;
        }}
        @media print {{
            body {{
                background-color: white;
            }}
            .container {{
                box-shadow: none;
                padding: 0;
            }}
        }}
    </style>
</head>
<body>
    <div class="container">
        <h1>{metadata.get('name', 'Briefing')}</h1>
        <div class="metadata">
            Generated: {metadata.get('generated_at', datetime.now().isoformat())}
        </div>
        {html_body}
        <div class="footer">
            This briefing contains {metadata.get('item_count', 0)} items from {metadata.get('source_count', 0)} sources.
        </div>
    </div>
</body>
</html>"""
            
            # Write HTML file
            output_path.write_text(html, encoding='utf-8')
            
            logger.info(f"Exported HTML to: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error exporting HTML: {str(e)}")
            raise
    
    def _convert_markdown_for_pdf(self, text: str) -> str:
        """Convert markdown formatting to ReportLab XML tags."""
        import re
        
        # Bold
        text = re.sub(r'\*\*([^*]+)\*\*', r'<b>\1</b>', text)
        text = re.sub(r'__([^_]+)__', r'<b>\1</b>', text)
        
        # Italic
        text = re.sub(r'\*([^*]+)\*', r'<i>\1</i>', text)
        text = re.sub(r'_([^_]+)_', r'<i>\1</i>', text)
        
        # Links (simplified - just show text)
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
        
        # Inline code
        text = re.sub(r'`([^`]+)`', r'<font name="Courier">\1</font>', text)
        
        return text
    
    def _add_markdown_to_paragraph(self, paragraph, text: str):
        """Add markdown-formatted text to DOCX paragraph."""
        import re
        
        # Split text by formatting markers
        parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|__[^_]+__|_[^_]+_|`[^`]+`)', text)
        
        for part in parts:
            if not part:
                continue
            
            # Bold
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('__') and part.endswith('__'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            
            # Italic
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('_') and part.endswith('_') and not part.startswith('__'):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            
            # Code
            elif part.startswith('`') and part.endswith('`'):
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Courier New'
            
            # Regular text
            else:
                paragraph.add_run(part)


# End of export_manager.py